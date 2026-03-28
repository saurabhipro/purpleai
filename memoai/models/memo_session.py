# -*- coding: utf-8 -*-
import base64
import io
import json
import logging
from datetime import datetime
from odoo import models, fields, api, _
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)


class MemoSession(models.Model):
    """
    Analysis Session — the main workflow record.
    Tracks the user's 5-step AI analysis journey for a selected subject.
    """
    _name = 'memo_ai.session'
    _description = 'Memo AI Analysis Session'
    _inherit = ['mail.thread']
    _order = 'create_date desc'

    name = fields.Char(string='Session Reference', readonly=True, copy=False, default='New')
    subject_id = fields.Many2one('memo_ai.subject', string='Subject', required=True,
                                 tracking=True, ondelete='restrict')
    company_id = fields.Many2one('res.company', default=lambda s: s.env.company)
    user_id = fields.Many2one('res.users', string='Analyst', default=lambda s: s.env.user, tracking=True)
    create_date = fields.Datetime(string='Created On', readonly=True)

    state = fields.Selection([
        ('draft', 'Draft'),
        ('step1_done', 'Step 1 Done'),
        ('step2_done', 'Step 2 Done'),
        ('step3_done', 'Step 3 Done'),
        ('step4_done', 'Step 4 Done'),
        ('done', 'Complete'),
    ], default='draft', string='Status', tracking=False)

    # ── Uploaded Documents ──────────────────────────────────────────────────────
    document_ids = fields.Many2many(
        'ir.attachment', 'memo_session_attachment_rel',
        'session_id', 'attachment_id',
        string='Source Documents',
        help="Upload PDFs or scanned documents for analysis."
    )

    # ── Step 1: Summary ─────────────────────────────────────────────────────────
    step1_output = fields.Html(string='Step 1 — Summary', sanitize=False)
    step1_processing = fields.Boolean(default=False)
    step1_last_run = fields.Datetime(string='Step 1 Last Run')
    step1_iteration = fields.Integer(string='Step 1 Iterations', default=0)
    step1_sources = fields.Text(string='Step 1 Source Documents (RAG)')

    # ── Step 2: Issue Identification ────────────────────────────────────────────
    step2_output = fields.Html(string='Step 2 — Applicable Issues', sanitize=False)
    step2_processing = fields.Boolean(default=False)
    step2_last_run = fields.Datetime(string='Step 2 Last Run')
    step2_iteration = fields.Integer(string='Step 2 Iterations', default=0)
    step2_sources = fields.Text(string='Step 2 Source Documents (RAG)')

    # ── Step 3: Regulatory Guidelines ──────────────────────────────────────────
    step3_output = fields.Html(string='Step 3 — Regulatory Guidelines', sanitize=False)
    step3_processing = fields.Boolean(default=False)
    step3_last_run = fields.Datetime(string='Step 3 Last Run')
    step3_iteration = fields.Integer(string='Step 3 Iterations', default=0)
    step3_sources = fields.Text(string='Step 3 Source Documents (RAG)')

    # ── Step 4: Analysis ────────────────────────────────────────────────────────
    step4_output = fields.Html(string='Step 4 — Analysis', sanitize=False)
    step4_processing = fields.Boolean(default=False)
    step4_last_run = fields.Datetime(string='Step 4 Last Run')
    step4_iteration = fields.Integer(string='Step 4 Iterations', default=0)
    step4_sources = fields.Text(string='Step 4 Source Documents (RAG)')

    # ── Chat Features (Split Section Chat) ──────────────────────────────────────
    step1_chat_enabled = fields.Boolean(string='Step 1 Chat Enabled', default=False)
    step1_chat_history = fields.Json(string='Step 1 Chat History', default=[])
    step1_chat_input = fields.Char(string='Step 1 Chat Input') # For simple form entry

    step2_chat_enabled = fields.Boolean(string='Step 2 Chat Enabled', default=False)
    step2_chat_history = fields.Json(string='Step 2 Chat History', default=[])
    step2_chat_input = fields.Char(string='Step 2 Chat Input')

    step3_chat_enabled = fields.Boolean(string='Step 3 Chat Enabled', default=False)
    step3_chat_history = fields.Json(string='Step 3 Chat History', default=[])
    step3_chat_input = fields.Char(string='Step 3 Chat Input')

    step4_chat_enabled = fields.Boolean(string='Step 4 Chat Enabled', default=False)
    step4_chat_history = fields.Json(string='Step 4 Chat History', default=[])
    step4_chat_input = fields.Char(string='Step 4 Chat Input')

    # ── Computed Chat UI Displays ──────────────────────────────────────────────
    step1_chat_display = fields.Html(compute='_compute_chat_displays')
    step2_chat_display = fields.Html(compute='_compute_chat_displays')
    step3_chat_display = fields.Html(compute='_compute_chat_displays')
    step4_chat_display = fields.Html(compute='_compute_chat_displays')

    @api.depends('step1_chat_history', 'step2_chat_history', 'step3_chat_history', 'step4_chat_history')
    def _compute_chat_displays(self):
        """Pre-render the chat bubbles for each step to avoid OWL directive errors in form views."""
        for rec in self:
            for s in [1, 2, 3, 4]:
                history = getattr(rec, f'step{s}_chat_history') or []
                color_map = {1: '#007bff', 2: '#ffc107', 3: '#17a2b8', 4: '#198754'}
                text_color_map = {1: '#ffffff', 2: '#212529', 3: '#ffffff', 4: '#ffffff'}
                
                if not history:
                    html = f"""
                        <div class="text-center text-muted mt-5 opacity-50">
                            <i class="fa fa-comments-o fa-3x mb-3"/>
                            <p>Ask questions about Step {s} context.</p>
                        </div>
                    """
                else:
                    parts = []
                    for msg in history:
                        role = msg.get('role', 'user')
                        content = msg.get('content', '')
                        is_user = role == 'user'
                        bg = color_map[s] if is_user else '#ffffff'
                        tc = text_color_map[s] if is_user else '#212529'
                        justify = 'justify-content-end' if is_user else 'justify-content-start'
                        border = 'border' if not is_user else ''
                        
                        bubble = f"""
                            <div class="mb-3 d-flex {justify}">
                                <div class="px-3 py-2 rounded-3 shadow-sm {border}" 
                                     style="max-width: 85%; line-height: 1.4; background-color: {bg}; color: {tc};">
                                    {content}
                                </div>
                            </div>
                        """
                        parts.append(bubble)
                    html = "".join(parts)
                
                setattr(rec, f'step{s}_chat_display', html)

    # ── Tracking & Analytics ────────────────────────────────────────────────────
    total_time_seconds = fields.Float(string='Time Taken (s)', default=0.0,
                                      help="Total execution time for all AI calls in this session.")
    total_cost = fields.Float(string='API Cost ($)', digits=(10, 4), default=0.0,
                              help="Estimated USD cost based on token consumption.")

    # ── Step 5: Word Export ─────────────────────────────────────────────────────
    word_attachment_id = fields.Many2one('ir.attachment', string='Word Document', readonly=True)

    @api.model_create_multi
    def create(self, vals_list):
        for vals in vals_list:
            if vals.get('name', 'New') == 'New':
                vals['name'] = self.env['ir.sequence'].next_by_code('memo_ai.session') or 'New'
        return super().create(vals_list)

    def write(self, vals):
        """Custom logging for HTML steps since native Odoo tracking doesn't support HTML types."""
        res = super(MemoSession, self).write(vals)
        for field in ['step1_output', 'step2_output', 'step3_output', 'step4_output']:
            if field in vals and not self.env.context.get('ai_processing'):
                # Very concise log for manual overrides
                self.message_post(body=_("Edited %s") % field.replace('_output', '').replace('step', 'Step '))
        return res

    def _update_state_safely(self, target_state, vals):
        """Helper to update state only if it moves forward in the workflow."""
        self.ensure_one()
        STATES = ['draft', 'step1_done', 'step2_done', 'step3_done', 'step4_done', 'done']
        
        def get_rank(s):
            try: return STATES.index(s or 'draft')
            except: return 0
            
        # Merge the new values into a dict
        vals = dict(vals)
        if get_rank(target_state) > get_rank(self.state):
            vals['state'] = target_state
        
        # Reset processing states
        vals['step1_processing'] = False
        vals['step2_processing'] = False
        vals['step3_processing'] = False
        vals['step4_processing'] = False
        
        return self.with_context(ai_processing=True).write(vals)

    # ───────────────────────────────────────────────────────────────────────────
    # Step 1: Extract & Summarize
    # ───────────────────────────────────────────────────────────────────────────
    def action_run_step1(self):
        """Extract a summary from uploaded PDFs using LLM."""
        self.ensure_one()
        self.step1_processing = True

        try:
            if not self.document_ids:
                raise UserError(_("Please upload at least one document before running Step 1."))
            subject = self.subject_id
            if not subject.summarization_prompt:
                raise UserError(_("The subject '%s' has no Summarization Prompt configured.") % subject.name)

            # Extract text from all uploaded PDFs
            all_text_parts = []
            for attachment in self.document_ids:
                try:
                    file_data = base64.b64decode(attachment.datas)
                    rag_doc = self.env['memo_ai.rag_document'].new({})
                    text = rag_doc._extract_pdf_text(file_data, attachment.name)
                    all_text_parts.append(f"=== {attachment.name} ===\n{text}")
                except Exception as e:
                    _logger.warning("Could not extract text from %s: %s", attachment.name, str(e))
                    all_text_parts.append(f"=== {attachment.name} ===\n[Could not extract text: {e}]")

            combined_text = "\n\n".join(all_text_parts)

            # Build prompt
            prompt = subject.summarization_prompt.replace('{document_text}', combined_text)

            # Call AI
            ai_data = self._call_ai(prompt)
            doc_names = ", ".join(self.document_ids.mapped('name'))
            self._update_state_safely('step1_done', {
                'step1_output': ai_data.get('text', ''),
                'step1_last_run': fields.Datetime.now(),
                'step1_iteration': self.step1_iteration + 1,
                'step1_sources': doc_names,
                'total_time_seconds': self.total_time_seconds + ai_data.get('elapsed', 0),
                'total_cost': self.total_cost + ai_data.get('cost', 0.0),
            })
            self.message_post(body=_("✅ Step 1 (Summary) completed by AI."))
        finally:
            self.write({'step1_processing': False})
        return True

    def action_step1_toggle_chat(self):
        self.ensure_one()
        self.step1_chat_enabled = not self.step1_chat_enabled

    def action_chat_step(self):
        """Unified chat logic for any step using its specific RAG context from self.env.context."""
        self.ensure_one()
        step_num = self.env.context.get('step_num')
        if not step_num or step_num not in [1, 2, 3, 4]:
            msg = _("Invalid Step selected for chat.")
            raise UserError(msg)
            
        input_field = f'step{step_num}_chat_input'
        history_field = f'step{step_num}_chat_history'
        output_field = f'step{step_num}_output'
        rag_type_map = {1: 'analysis', 2: 'issue_list', 3: 'guideline', 4: 'analysis'}
        
        from odoo.tools import html2plaintext
        user_message_raw = getattr(self, input_field)
        if not user_message_raw:
            return True 
            
        user_message = html2plaintext(user_message_raw).strip()
        if not user_message or len(user_message) < 2:
            return True # Too short or empty
            
        # Get existing history
        history = list(getattr(self, history_field) or [])
        history.append({'role': 'user', 'content': user_message})
        
        # Build Context: Current Step Output + RAG
        # Performance: We only embed the USER MESSAGE for searching, not the whole summary
        rag_matches = self.env['memo_ai.rag_document'].search_vector_similarity(
            user_message, 
            limit=3, 
            subject_type=rag_type_map.get(step_num, 'analysis')
        )
        rag_text = "\n\n".join([m['content'] for m in rag_matches])
        
        # Build the Prompt
        full_prompt = f"""
            You are a subject-matter expert chatting about Step {step_num} of this analysis for the subject: {self.subject_id.name}.
            
            REFERENCE CONTEXT (CURRENT ANALYSIS - STEP {step_num}):
            {(getattr(self, output_field) or "")[:5000]}
            
            ADDITIONAL RAG CONTEXT:
            {rag_text}
            
            CHAT HISTORY:
            {json.dumps(history[-4:], indent=2)}
            
            USER MESSAGE:
            {user_message}
            
            INSTRUCTIONS:
            1. Answer the user's question precisely using the contexts above.
            2. If the user asks for changes to the analysis, explain how they should edit the HTML box.
            3. Keep responses helpful but concise.
        """
        
        ai_data = self._call_ai(full_prompt)
        ai_response = ai_data.get('text', '')
        
        history.append({'role': 'assistant', 'content': ai_response})
        
        self.write({
            history_field: history,
            input_field: False,
            'total_time_seconds': self.total_time_seconds + ai_data.get('elapsed', 0),
            'total_cost': self.total_cost + ai_data.get('cost', 0.0),
        })
        return True

    def action_rerun_step1(self):
        """Allow manual rerun of Step 1."""
        self.ensure_one()
        return self.action_run_step1()

    # ───────────────────────────────────────────────────────────────────────────
    # Step 2: Issue Identification via RAG
    # ───────────────────────────────────────────────────────────────────────────
    def action_run_step2(self):
        """Search issue list RAG and identify applicable issues."""
        self.ensure_one()
        if not self.step1_output:
            raise UserError(_("Please complete Step 1 (Summary) first."))
        
        self.step2_processing = True

        try:
            subject = self.subject_id
            if not subject.issue_extraction_prompt:
                raise UserError(_("The subject '%s' has no Issue Extraction Prompt configured.") % subject.name)

            # Retrieve RAG context from issue list documents
            prompt_context = self.step1_output
            rag_matches = self.env['memo_ai.rag_document'].search_vector_similarity(
                prompt_context, 
                limit=8, 
                subject_type='issue_list'
            )
            rag_text = "\n\n".join([m['content'] for m in rag_matches])
            source_names = ", ".join(list(set([m['document_name'] for m in rag_matches if m.get('document_name')])))

            ai_prompt = subject.issue_extraction_prompt\
                .replace('{summary}', self.step1_output or "")\
                .replace('{rag_context}', rag_text or "No issue list RAG documents found.")

            ai_data = self._call_ai(ai_prompt)
            
            self._update_state_safely('step2_done', {
                'step2_output': ai_data.get('text', ''),
                'step2_last_run': fields.Datetime.now(),
                'step2_iteration': self.step2_iteration + 1,
                'step2_sources': source_names or "Internal Knowledge",
                'total_time_seconds': self.total_time_seconds + ai_data.get('elapsed', 0),
                'total_cost': self.total_cost + ai_data.get('cost', 0.0),
            })
            self.message_post(body=_("✅ Step 2 (Issue Identification) completed by AI."))
        finally:
            self.write({'step2_processing': False})
        return True

    def action_step2_toggle_chat(self):
        self.ensure_one()
        self.step2_chat_enabled = not self.step2_chat_enabled

    def action_rerun_step2(self):
        """Allow manual rerun of Step 2."""
        self.ensure_one()
        return self.action_run_step2()

    # ───────────────────────────────────────────────────────────────────────────
    # Step 3: Regulatory Guidelines
    # ───────────────────────────────────────────────────────────────────────────
    def action_run_step3(self):
        """Find applicable regulatory guidelines using guideline RAG."""
        self.ensure_one()
        if not self.step2_output:
            raise UserError(_("Please complete Step 2 (Issues) first."))
        
        self.step3_processing = True

        try:
            subject = self.subject_id
            if not subject.regulatory_extraction_prompt:
                raise UserError(_("The subject '%s' has no Regulatory Extraction Prompt configured.") % subject.name)

            query = (self.step2_output or "") + " " + (self.step1_output or "")
            rag_matches = self.env['memo_ai.rag_document'].search_vector_similarity(
                query, 
                limit=8, 
                subject_type='guideline'
            )
            rag_text = "\n\n".join([m['content'] for m in rag_matches])
            source_names = ", ".join(list(set([m['document_name'] for m in rag_matches if m.get('document_name')])))

            prompt = subject.regulatory_extraction_prompt\
                .replace('{summary}', self.step1_output or "")\
                .replace('{issues}', self.step2_output or "")\
                .replace('{rag_context}', rag_text or "No guideline RAG documents found.")

            ai_data = self._call_ai(prompt)
            self._update_state_safely('step3_done', {
                'step3_output': ai_data['text'],
                'step3_last_run': fields.Datetime.now(),
                'step3_iteration': self.step3_iteration + 1,
                'step3_sources': source_names or "Internal Knowledge",
                'total_time_seconds': self.total_time_seconds + ai_data['elapsed'],
                'total_cost': self.total_cost + ai_data['cost'],
            })
            self.message_post(body=_("✅ Step 3 (Regulatory Guidelines) completed by AI."))
        finally:
            self.write({'step3_processing': False})
        return True

    def action_step3_toggle_chat(self):
        self.ensure_one()
        self.step3_chat_enabled = not self.step3_chat_enabled

    def action_rerun_step3(self):
        """Allow manual rerun of Step 3."""
        self.ensure_one()
        return self.action_run_step3()

    # ───────────────────────────────────────────────────────────────────────────
    # Step 4: Final Analysis
    # ───────────────────────────────────────────────────────────────────────────
    def action_run_step4(self):
        """Run analysis combining all previous steps + Analysis RAG."""
        self.ensure_one()
        if not self.step3_output:
            raise UserError(_("Please complete Step 3 (Regulations) first."))
        
        self.step4_processing = True

        try:
            subject = self.subject_id
            if not subject.analysis_prompt:
                raise UserError(_("The subject '%s' has no Analysis Prompt configured.") % subject.name)

            query = " ".join(filter(None, [self.step1_output, self.step2_output, self.step3_output]))
            rag_matches = self.env['memo_ai.rag_document'].search_vector_similarity(
                query, 
                limit=8, 
                subject_type='analysis'
            )
            rag_text = "\n\n".join([m['content'] for m in rag_matches])
            source_names = ", ".join(list(set([m['document_name'] for m in rag_matches if m.get('document_name')])))

            ai_prompt = subject.analysis_prompt\
                .replace('{summary}', self.step1_output or "")\
                .replace('{issues}', self.step2_output or "")\
                .replace('{regulatory}', self.step3_output or "")\
                .replace('{rag_context}', rag_text or "No analysis RAG documents found.")

            ai_data = self._call_ai(ai_prompt)
            
            self._update_state_safely('step4_done', {
                'step4_output': ai_data.get('text', ''),
                'step4_last_run': fields.Datetime.now(),
                'step4_iteration': self.step4_iteration + 1,
                'step4_sources': source_names or "Synthesized Highlights",
                'total_time_seconds': self.total_time_seconds + ai_data.get('elapsed', 0),
                'total_cost': self.total_cost + ai_data.get('cost', 0.0),
            })
            self.message_post(body=_("✅ Step 4 (Analysis) completed by AI."))
        finally:
            self.write({'step4_processing': False})
        return True

    def action_step4_toggle_chat(self):
        self.ensure_one()
        self.step4_chat_enabled = not self.step4_chat_enabled

    def action_rerun_step4(self):
        """Allow manual rerun of Step 4."""
        self.ensure_one()
        return self.action_run_step4()

    # ───────────────────────────────────────────────────────────────────────────
    # Step 5: Export to Word
    # ───────────────────────────────────────────────────────────────────────────
    # ───────────────────────────────────────────────────────────────────────────
    # ── Export & Export Wizard ─────────────────────────────────────────────────
    # ───────────────────────────────────────────────────────────────────────────
    def action_open_export_wizard(self):
        """Show a dialog for PDF/Word format choice."""
        self.ensure_one()
        return {
            'name': _('Export Analysis Report'),
            'type': 'ir.actions.act_window',
            'res_model': 'memo_ai.export_wizard',
            'view_mode': 'form',
            'target': 'new',
            'context': {'default_session_id': self.id}
        }

    def action_export_pdf(self):
        # Placeholder for PDF export (will create Word first then maybe convert or handle)
        return self.action_export_word()

    def action_export_word(self):
        """Generate a Word (.docx) document from all step outputs."""
        self.ensure_one()
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise UserError(_(
                "python-docx is not installed. Please run: pip install python-docx"
            ))

        doc = Document()

        # Title
        title = doc.add_heading(f"Analysis Memo — {self.subject_id.name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Reference: {self.name}")
        doc.add_paragraph(f"Analyst: {self.user_id.name}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
        doc.add_paragraph()

        sections = [
            ("1. Summary", self.step1_output),
            ("2. Applicable Issues", self.step2_output),
            ("3. Regulatory Guidelines", self.step3_output),
            ("4. Analysis", self.step4_output),
        ]
        from odoo.tools import html2plaintext
        for heading, content in sections:
            if content:
                doc.add_heading(heading, level=1)
                doc.add_paragraph(html2plaintext(content))
                doc.add_paragraph()

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        file_data = base64.b64encode(buffer.read()).decode('utf-8')

        attachment = self.env['ir.attachment'].create({
            'name': f"MemoAI_{self.name}_{self.subject_id.name}.docx",
            'type': 'binary',
            'datas': file_data,
            'res_model': self._name,
            'res_id': self.id,
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        })

        self.write({
            'word_attachment_id': attachment.id,
            'state': 'done',
        })
        # Removed message_post for Word generation to keep chatter clean

        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'self',
        }


    # ───────────────────────────────────────────────────────────────────────────
    # ── AI Call Helper ──────────────────────────────────────────────────────────
    # ───────────────────────────────────────────────────────────────────────────
    def _call_ai(self, prompt):
        """
        Call the configured AI provider. Returns a dict with {text, cost, elapsed}.
        """
        import time
        try:
            from ..services.memo_ai_service import call_ai
            
            start_t = time.time()
            ai_data = call_ai(self.env, prompt)
            elapsed = time.time() - start_t
            
            if isinstance(ai_data, dict):
                return {
                    'text': ai_data.get('text', ''),
                    'cost': ai_data.get('cost', 0.0),
                    'elapsed': elapsed
                }
            else:
                return {
                    'text': ai_data or '',
                    'cost': 0.0,
                    'elapsed': elapsed
                }
            
        except Exception as e:
            _logger.error("AI call failed: %s", str(e))
            raise UserError(_(f"AI processing failed: {e}\n\nCheck Memo AI → Configuration → Settings."))

    def action_reset_to_draft(self):
        """Reset session to draft (clears all step outputs)."""
        self.write({
            'state': 'draft',
            'step1_output': False,
            'step1_last_run': False,
            'step1_iteration': 0,
            'step2_output': False,
            'step2_last_run': False,
            'step2_iteration': 0,
            'step3_output': False,
            'step3_last_run': False,
            'step3_iteration': 0,
            'step4_output': False,
            'step4_last_run': False,
            'step4_iteration': 0,
            'total_cost': 0.0,
            'total_time_seconds': 0.0,
            'step1_chat_enabled': False,
            'step2_chat_enabled': False,
            'step3_chat_enabled': False,
            'step4_chat_enabled': False,
            'step1_chat_history': [],
            'step2_chat_history': [],
            'step3_chat_history': [],
            'step4_chat_history': [],
        })

    @api.model
    def get_dashboard_stats(self, date_from=None, date_to=None):
        """Native RPC method to fetch data for the AI Analytics Dashboard Component."""
        domain = []
        if date_from:
            domain.append(('create_date', '>=', date_from))
        if date_to:
            domain.append(('create_date', '<=', date_to))
            
        sessions = self.search(domain)
        
        # Aggregate Top KPIs
        total_requests = len(sessions)
        passes = len(sessions.filtered(lambda s: s.state == 'done'))
        # Considered failed/in process if not done and not draft
        in_process = len(sessions.filtered(lambda s: s.state not in ('done', 'draft')))
        total_time = sum(sessions.mapped('total_time_seconds'))
        total_cost = sum(sessions.mapped('total_cost'))
        
        # Aggregate Chart Data for Bar Chart: cost and time group by date
        chart_dict = {}
        for s in sessions:
            if not s.create_date:
                continue
            day_str = s.create_date.strftime('%Y-%m-%d')
            if day_str not in chart_dict:
                chart_dict[day_str] = {'cost': 0.0, 'time': 0.0}
            chart_dict[day_str]['cost'] += s.total_cost
            chart_dict[day_str]['time'] += s.total_time_seconds
            
        labels = sorted(list(chart_dict.keys()))
        cost_data = [chart_dict[l]['cost'] for l in labels]
        time_data = [chart_dict[l]['time'] for l in labels]
            
        return {
            'total_requests': total_requests,
            'passes': passes,
            'in_process': in_process,
            'total_time': round(total_time, 2),
            'total_cost': round(total_cost, 4),
            'chart': {
                'labels': labels,
                'cost_data': cost_data,
                'time_data': time_data
            }
        }

class MemoExportWizard(models.TransientModel):
    _name = 'memo_ai.export_wizard'
    _description = 'Memo Export Wizard'

    session_id = fields.Many2one('memo_ai.session', string='Session', required=True)
    format = fields.Selection([
        ('word', 'Word (.docx)'),
        ('pdf', 'Portable Document (.pdf) — COMING SOON')
    ], string='Export Format', default='word', required=True)

    def action_export(self):
        self.ensure_one()
        if self.format == 'word':
            return self.session_id.action_export_word()
        else:
            return self.session_id.action_export_pdf()
